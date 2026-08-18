"""
Build Healthcare_Secure_Handbook.docx — single source of truth for students.
No real secrets, IPs with credentials, or private keys are included.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "handbook_assets"
OUT = ROOT / "Healthcare_Secure_Handbook.docx"
OUT_FALLBACK = ROOT / "Healthcare_Secure_Handbook_26Jul2026.docx"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True)
    return h


def add_p(doc, text, *, bold=False, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        for run in p.runs:
            set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            set_run_font(run, size=11)


def add_image(doc, name, caption, width=6.2):
    path = ASSETS / name
    if not path.exists():
        add_p(doc, f"[Missing image: {name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            for p in cells[c_i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ----- Cover -----
    add_image(doc, "00_cover_art.png", "Cover illustration (decorative)", width=6.0)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Aarogya One Connect")
    set_run_font(r, size=28, bold=True, color=RGBColor(0x1A, 0x52, 0x76))
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("Healthcare Secure — Complete Student Handbook & Single Source of Truth")
    set_run_font(r, size=14, bold=True, color=RGBColor(0x1A, 0x52, 0x76))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "From first scaffold to current clinic app — architecture, Android/Capacitor, "
        "analytics, ABDM/SMS/multi-tenant, gynae decision support, cloud Hostinger ops, "
        "patient billing / Razorpay UPI QR, video consult, browser desk (app.*), "
        "Today-only patient pick, persisted lab orders, stale-APK + CORS RCA, secrets, VPS deploy"
    )
    set_run_font(r, size=12)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Application version baseline: 0.1.0 (+ Unreleased features documented herein)\n"
        "Document edition: 15 August 2026 (includes Era O: Today-only pick, doctor→Visit after lock, "
        "persisted lab orders, 15s chart poll, browser desk at app.*, CORS Failed-to-fetch RCA, "
        "Lab Desk 403 lock RCA; share APK v1.48)\n"
        "Audience: students learning Android (Capacitor), FastAPI, and clinic workflows\n\n"
        "SECURITY NOTICE: This document never contains real SECRET_KEY, SECRET_SALT, "
        "PINs, API keys, Razorpay secrets, or production passwords. Use placeholders only."
    )
    set_run_font(r, size=9, italic=True, color=RGBColor(0x92, 0x2B, 0x21))
    page_break(doc)

    # ----- TOC -----
    add_heading(doc, "Table of contents", 1)
    toc_items = [
        "Part I — Foundations",
        "Part II — Learning paths (Android, patient analytics, backend)",
        "Part III — Chronological development (scratch → now)",
        "Part IV — System deep dive (how it works today)",
        "Part V — Issues, debugging strategies, and root cause analysis",
        "Part VI — Secrets, changing app behavior, and paid VPS deployment",
        "Part VII — Future enhancement plans",
        "Part VIII — Appendices (glossary, commands, file tree)",
    ]
    add_numbered(doc, toc_items)
    add_p(
        doc,
        "This Word document is the intended single source of truth. "
        "README.md in the repo is a short entry point that points here. "
        "CHANGELOG.md remains the SemVer history; annexes: docs/PLAY_STORE.md, "
        "docs/CLOUD_DEPLOY.md, docs/RESTART_PRODUCTION_API.md, docs/PRIVACY_POLICY.md, "
        "docs/ABDM_SMS_MULTI_TENANT.md, docs/DEMO_CLIENT.md. Browser desk: https://app.aarogyaoneconnect.in.",
    )
    page_break(doc)

    # ========== PART I ==========
    add_heading(doc, "Part I — Foundations", 1)
    add_heading(doc, "1. Problem statement", 2)
    add_p(
        doc,
        "Small clinics often juggle paper prescriptions, scattered phone photos of lab reports, "
        "and WhatsApp forwards. That is hard to search, easy to lose, and risky for privacy. "
        "Healthcare Secure is a dual-root clinic application: a Python FastAPI backend that stores "
        "de-identified clinical encounters, and a Next.js + Capacitor frontend that runs in a browser "
        "or as an Android app on ordinary staff phones.",
    )
    add_p(
        doc,
        "Significance for learners: you are not building a toy todo list — you are practicing "
        "identity minimization, role-based access, mobile packaging, and operational deploy thinking "
        "on a domain that cares about patient data.",
    )

    add_heading(doc, "2. Product constraints", 2)
    add_bullets(
        doc,
        [
            "Non-technical staff UI — short labels, Hindi/English toggle, large tap targets.",
            "Roles: doctor (sign Rx), staff (vitals/notes), lab (results/uploads; Today/Labs/More only), "
            "receptionist (Patient Info + billing/appointments; no Visit/Rx). Same UI on Android APK "
            "and browser desk https://app.aarogyaoneconnect.in.",
            "Gynecology-oriented decision support (obstetric card, ANC alerts, case brief) is "
            "advisory only — never auto-orders or replaces clinical judgment.",
            "Phones are general devices carried everywhere — open app on mobile data and work "
            "(clinic PC is optional for day-to-day use).",
            "Security first: raw name/phone/MRN are HMAC-blinded before persistence; secrets stay out of git.",
        ],
    )

    add_heading(doc, "3. Security mindset for beginners", 2)
    add_p(
        doc,
        "PHI (Protected / personal health information) includes clinical notes, vitals, labs, and documents. "
        "Even if you hash the patient’s phone, the visit notes remain sensitive. Blind IDs reduce the risk "
        "of a stolen database leaking a phone-number column — they do not make the whole dataset “anonymous.”",
    )
    add_image(doc, "02_blind_identity.png", "Figure: HMAC blind identity pipeline")
    add_bullets(
        doc,
        [
            "SECRET_SALT — used only for HMAC patient IDs. Changing it remaps history (plan carefully).",
            "SECRET_KEY — used for signed prescription download URLs.",
            "PINs — prefer pbkdf2$… hashes from scripts/hash_pin.py; lockout after failed attempts.",
            "TLS — phones on public 4G must use HTTPS with a publicly trusted certificate (Let’s Encrypt), "
            "not a self-signed LAN cert.",
        ],
    )

    add_heading(doc, "4. Tech stack glossary", 2)
    add_table(
        doc,
        ["Term", "Plain meaning in this project"],
        [
            ["FastAPI", "Python web framework for the clinic API"],
            ["SQLAlchemy", "ORM mapping Python classes to DB tables"],
            ["SQLite / Postgres", "Local file DB vs production server DB"],
            ["Next.js App Router", "React UI framework; we static-export for Capacitor"],
            ["Capacitor", "Wraps web UI in a native Android shell"],
            ["Whisper", "Speech-to-text (local faster-whisper or cloud)"],
            ["Ollama / LLM", "Parses transcript text into structured Rx JSON"],
            ["HMAC", "Keyed hash — same salt + input → same blind ID"],
            ["nginx", "Reverse proxy terminating TLS in Docker"],
            ["Capacitor webDir", "frontend/out copied into the APK — UI updates need a new APK"],
            ["Razorpay UPI QR", "Staff shows QR; webhook posts payment to the patient ledger"],
            ["AAAA record", "DNS IPv6 address — needed when some ISPs black-hole VPS IPv4"],
            ["Browser desk", "Static Next export at https://app.* — same clinic UI as the APK"],
            ["lab_orders_json", "Visit diagnostic ticks stored on clinic_patients for Lab Desk"],
        ],
    )

    add_heading(doc, "5. Dual-root layout", 2)
    add_image(doc, "07_dual_root.png", "Figure: backend/ and frontend/ dual roots")
    add_p(
        doc,
        "Keeping API and UI in separate roots teaches a production habit: the phone never embeds the "
        "database; it only calls HTTPS endpoints. You can rebuild the Android app without touching "
        "Python, and scale the API independently later.",
    )
    page_break(doc)

    # ========== PART II ==========
    add_heading(doc, "Part II — Learning paths", 1)

    add_heading(doc, "6. Android path (Capacitor)", 2)
    add_p(
        doc,
        "The “Android app” is mostly your Next.js UI compiled to static files in frontend/out/, "
        "then synced into a native project with Capacitor. The WebView loads those files. Native "
        "plugins add microphone, share, and print behaviors.",
    )
    add_numbered(
        doc,
        [
            "Set frontend/.env.local NEXT_PUBLIC_API_BASE_URL (optional default; runtime override exists).",
            "npm run mobile:build → next build + npx cap sync.",
            "npm run cap:open:android → Android Studio.",
            "Run on USB device; grant RECORD_AUDIO when prompted.",
            "On sign-in, set Clinic server address to your API (LAN or https://api.example).",
        ],
    )
    add_p(
        doc,
        "Significance: static export (output: \"export\" in next.config) is required so there is no "
        "Node server on the phone. Default Capacitor builds allow cleartext HTTP for LAN learning. "
        "For cloud HTTPS production APKs set CAPACITOR_HTTPS=true before npm run mobile:build "
        "(or run scripts/build_share_apk.cmd from repo root — preferred share build). "
        "AGP 9.x needs proguard-android-optimize.txt — "
        "frontend/scripts/fix-agp-proguard.mjs runs during cap:sync / mobile:build. "
        "Android versionCode/versionName auto-bump on assemble (frontend/android/app/version.properties).",
    )
    add_p(
        doc,
        "Critical: the phone does NOT load UI from the cloud API. Deploying or restarting the VPS "
        "API only changes backend behavior. After any frontend change you must next build → "
        "cap sync → assemble APK → reinstall on phones. More → “UI build: …” shows which "
        "static bundle is inside the installed APK (see RCA-22).",
        bold=True,
    )
    add_p(
        doc,
        "Student tip: frontend/android/ is gitignored — each machine generates it with Capacitor. "
        "That surprises beginners who clone the repo and expect Android Studio files to appear automatically.",
    )

    add_heading(doc, "7. Patient analytics path (as implemented today)", 2)
    add_p(
        doc,
        "Clinic Analytics is live under More (not a hospital BI warehouse). Staff/doctors see "
        "operational counts; lab sees a limited lab-results summary. Patient Info (nav Today) "
        "holds All patients (the only name/MRN picker), waiting queue, appointments, and the "
        "patient billing ledger when a patient is locked (never shown to lab). On Visit, sections "
        "default collapsed — order is vitals trend → case brief → video consult (if enabled) → "
        "recommended diagnostics checklist → prescription / forward history. Doctor PIN/home lands "
        "on Patient Info to pick first; lock from All patients or waiting list opens Visit."
    )
    add_bullets(
        doc,
        [
            "GET /api/v1/analytics/today — encounters, unique patients, queue waiting, vitals, labs, Rx signed.",
            "GET /api/v1/analytics/week — last 7 days by day.",
            "GET /api/v1/analytics/frequency — top medications and diagnoses (30d).",
            "POST /api/v1/analytics/vitals-trend — weight/BP/Hb (+ pulse/SpO2/temp) over visits; merges "
            "lab Hb into the series; returns ANC rule alerts; optional gestational-age X-axis from LMP.",
            "GET /api/v1/analytics/export.csv — de-identified CSV (doctor only); blind-id prefix only.",
            "GET /api/v1/history/case-summary — consolidated obstetric + vitals deltas + labs + docs + "
            "scan cadence + soft alerts (decision support).",
            "POST /api/v1/history/consult-pack — LLM or rule-fallback “today’s consult” checklist.",
            "POST /api/v1/history/rx-hints — soft pre-sign medication conflict hints.",
            "POST /api/v1/history/documents/{id}/analyze — extract USG/report findings into structured JSON.",
            "Patient timeline, activity audit, waiting queue (Open → lock patient), structured lab + HL7 ORU.",
            "POST /api/v1/history/billing + billing-summary — charge/payment ledger; today’s charges, "
            "total paid, amount due (doctor/staff/receptionist).",
            "POST /api/v1/payments/qr — Razorpay UPI QR; poll status; webhook appends payment rows.",
            "Video consult (Jitsi): mint room + SMS join link from Visit (feature flag video_consult).",
            "GET/PUT /api/v1/history/lab-orders — persist Visit diagnostic ticks (doctor/staff/lab).",
            "Directory + locked chart poll ~15s so Lab Desk on another phone sees Visit orders.",
        ],
    )
    add_p(
        doc,
        "Student rule: never export raw phone/name columns. Temperature may be entered in °C or °F "
        "(stored as °F). Neonatal vitals bands apply when age is under ~28 days. LLM/consult output is "
        "decision support only — confirm before acting.",
    )

    add_heading(doc, "8. Backend path", 2)
    add_p(
        doc,
        "Request flow: HTTP → FastAPI router (app/api/v1/endpoints/*) → service (app/services/*) → "
        "SQLAlchemy model → database. Settings load from backend/.env via pydantic-settings "
        "(app/core/config.py). OpenAPI docs appear only when DEBUG=true.",
    )
    page_break(doc)

    # ========== PART III ==========
    add_heading(doc, "Part III — Chronological development", 1)
    add_p(
        doc,
        "This section is an engineering journal: each era shows goal, steps, significance, and how to verify. "
        "Dates refer to the July 2026 build history of this repository.",
    )

    milestones = [
        (
            "Era A — Scaffold (Day 0)",
            "Create dual-root FastAPI + Next.js App Router + Tailwind TypeScript skeleton with CORS locked to localhost and env-driven config.",
            [
                "backend/app/main.py with /health",
                "backend/app/core/config.py + .env.example",
                "frontend app router layout/page + Tailwind",
            ],
            "Teaches separation of API and UI and “no secrets in code.”",
            "curl http://127.0.0.1:8000/health ; npm run dev on :3000",
        ),
        (
            "Era B — Config & frontend versioning fixes",
            "Make uvicorn and Next actually start on the developer machine.",
            [
                "Fix CSV list env parsing with NoDecode",
                "Restore Next.js 15 after accidental downgrade to Next 9",
            ],
            "First real RCA lessons (see Part V).",
            "Uvicorn starts; Next shows the dashboard shell",
        ),
        (
            "Era C — Blind identity + clinical records",
            "Store encounters keyed by HMAC blind_patient_id, never raw phone columns.",
            [
                "app/services/security.py",
                "app/models/record.py ClinicalRecord",
                "SECRET_SALT in settings",
            ],
            "Core privacy architecture of the product.",
            "Tokenize endpoint returns shape mrn|… or name|phone without persisting raw IDs as columns",
        ),
        (
            "Era D — Voice → transcript → LLM → PDF",
            "Doctor dictates; system produces a signed prescription PDF and history row.",
            [
                "Whisper transcription service",
                "LLM parser (Ollama default)",
                "fpdf2 letterhead PDF",
                "Frontend VoiceRecorder + EncounterWorkspace",
            ],
            "End-to-end clinical value; introduces AI ops (CPU Whisper, Ollama).",
            "Record short clip → parse → write PDF",
        ),
        (
            "Era E — Mic Permissions-Policy bug",
            "Browser recording failed with NotAllowedError despite user gesture.",
            ["next.config Permissions-Policy microphone=(self)"],
            "Security headers can break features — always test mic after header changes.",
            "Record button captures audio on localhost",
        ),
        (
            "Era F — Mobile / LAN",
            "Run on Android over clinic Wi-Fi.",
            [
                "Capacitor config + mobile:build",
                "uvicorn --host 0.0.0.0",
                "CORS + PUBLIC_API_BASE_URL for LAN IP",
            ],
            "Bridges web skills into Android packaging.",
            "Phone browser or APK reaches API",
        ),
        (
            "Era G — Roles, vitals, simplified UI",
            "Staff enter vitals; lab uploads; doctor signs.",
            [
                "doctor/staff/lab roles",
                "vitals validation (°F)",
                "simplified non-tech copy",
            ],
            "Clinic realism beyond a single-doctor demo.",
            "Staff PIN can save vitals; lab cannot sign Rx",
        ),
        (
            "Era H — Reliability (P1)",
            "Survive API restarts; stop baking LAN IP into every APK rebuild.",
            [
                "Runtime clinic server URL",
                "Durable clinic_sessions",
                "PIN hash + lockout",
                "activity log, pytest CI, SQLite backup scripts",
            ],
            "Moves from prototype to operable clinic tool.",
            "Restart API; session still valid; change URL without rebuild",
        ),
        (
            "Era I — Enhancements (P2/P3 + later stubs)",
            "MRN identity, structured lab, print-first Rx, queue, i18n, offline queue, Docker/TLS, AAB docs, ABHA/HL7/pediatric stubs.",
            [
                "queue.py, integrations.py, LabResultsForm, WaitingQueue",
                "docker-compose.yml + nginx",
                "docs/PLAY_STORE.md",
            ],
            "Broadens clinic ops and prepares cloud hosting.",
            "Add queue entry; enter lab panel; docker compose builds",
        ),
        (
            "Era J — Ops, analytics, hardening (22 Jul 2026)",
            "Keep the API responsive on phones; ship clinic analytics; harden production defaults.",
            [
                "WHISPER_PRELOAD=false by default; warm Whisper in background if enabled",
                "scripts/watch_api.cmd health watchdog",
                "Phone first-run blocks localhost (needsClinicUrlSetup)",
                "ClinicAnalytics UI + /analytics/today|week|frequency|vitals-trend|export.csv",
                "Attachments on disk (content_path); TrustedHost + unlock rate-limit",
                "°C/°F preference; neonatal vitals; HL7 MSA ACK; CAPACITOR_HTTPS; PRIVACY_POLICY.md",
                "AGP ProGuard fix script (proguard-android-optimize.txt)",
            ],
            "Turns “LAN demo that hangs” into an operable daily clinic tool.",
            "curl /health answers in <1s; phone unlocks with LAN IP; analytics panel loads",
        ),
        (
            "Era K — ABDM, SMS appointments, multi-tenant (22–23 Jul 2026)",
            "National-ready ABHA OTP flow (sandbox/mock), appointment SMS, multi-clinic isolation.",
            [
                "abdm_client.py + /integrations/abha/otp/* + callback; ABDM_MOCK OTP 123456",
                "appointments API + SMS_PROVIDER=console|msg91|twilio; encrypted phones",
                "CLINICS + clinic_id on users/sessions/records/queue/appointments",
                "docs/ABDM_SMS_MULTI_TENANT.md",
            ],
            "Moves beyond single-clinic POC toward multi-site and ABDM-capable installs.",
            "Book appointment (console SMS log); ABDM_MOCK link; two clinics cannot see each other’s records",
        ),
        (
            "Era L — OPD shell, gynae decision support, demo seed (24–26 Jul 2026)",
            "Ship a multi-tab clinic shell and pregnancy-aware decision support without inventing diagnoses.",
            [
                "Clinic routes: Today / Patient / Visit / Records / Labs / More + ClinicNav bottom bar",
                "Obstetric profile (LMP/EDD/GPLA/blood group/Rh/high-risk) + LockedPatientChip GA",
                "Hemoglobin on vitals; pregnancy-aware weight/BP/Hb charts (GA X-axis when LMP set)",
                "Case brief UI + GET /history/case-summary; ANC BP/Hb/weight alerts (anc_alerts.py)",
                "Merge lab Hb into vitals-trend; ANC scan cadence checklist (NT / anomaly / growth)",
                "POST /history/consult-pack (LLM + rule fallback); POST /history/rx-hints before sign",
                "POST /history/documents/{id}/analyze for USG/report findings; Analyze on attachments",
                "Waiting List Open → lock patient; Google Calendar / .ics; day-before SMS remind script",
                "scripts/seed_demo.cmd — eight showcase patients covering alerts, cadence, Rx hints",
                "Offline queue: only queue network/5xx failures; drop 400/422 invalid jobs on sync",
                "Fix VitalsTrendCharts TS syntax (Python * keyword-only args broke mobile:build)",
            ],
            "Turns the app into a teachable gynae OPD demo: charts, alerts, consult pack, and safe Rx hints.",
            "seed_demo --wipe; lock Ananya Reddy → Visit case brief shows alerts; Records → Analyze growth USG; "
            "npm run mobile:build succeeds",
        ),
        (
            "Era M — Cloud Hostinger ops, dual-stack, video, Visit UX (late Jul – early Aug 2026)",
            "Run one shared HTTPS API for all clinics on mobile data; harden reachability; add receptionist + video.",
            [
                "Live API https://api.aarogyaoneconnect.in on Hostinger Mumbai VPS (/root/Healthcare-Secure)",
                "docs/CLOUD_DEPLOY.md + docs/RESTART_PRODUCTION_API.md; restart: unless-stopped on Compose services",
                "DNS A + AAAA for api (IPv6 path when some India ISPs black-hole VPS IPv4)",
                "After docker compose force-recreate api → also restart nginx (stale Docker DNS upstream → 502)",
                "Receptionist role: Patient Info + More (billing/appointments); CLINIC_USERS role receptionist",
                "Nav rename: Today tab = Patient Info; Patient tab = Vitals",
                "Visit: all sections default collapsed; order vitals trend → case brief → video → Rx",
                "Video consult feature (Jitsi): VIDEO_CONSULT_PROVIDER + JITSI_BASE_URL; SMS join link",
                "VPS .env must be /root/Healthcare-Secure/backend/.env (not ~/backend/.env); Compose $$ escape "
                "for pbkdf2$ hashes in CLINIC_USERS / CLINICS",
            ],
            "Makes “phones everywhere” real and documents production footguns that look like app bugs.",
            "Mobile-data unlock; curl /health over IPv4 and IPv6; Visit video panel when feature enabled; "
            "receptionist PIN sees Patient Info without Visit",
        ),
        (
            "Era N — Patient billing ledger + Razorpay UPI QR + stale-APK fix (Aug 2026)",
            "Front-desk billing on Patient Info; optional UPI QR pay; teach why phones lag source UI.",
            [
                "ClinicalRecord billing rows kind=charge|payment; POST /history/billing + billing-summary",
                "PatientBilling.tsx on Patient Info (doctor/staff/receptionist when patient locked)",
                "payment_intents model; /api/v1/payments/* create QR, poll, webhook, mock-pay",
                "PAYMENTS_ENABLED + RAZORPAY_* (+ RAZORPAY_MOCK for local tests); CLOUD_DEPLOY Razorpay section",
                "UI build stamp on More (APP_BUILD_ID); scripts/build_share_apk.cmd → share/AarogyaOneConnect-v*.apk",
                "Proved Aug 1 APK assets lacked Show pay QR while source had billing — rebuild + reinstall required",
            ],
            "Closes the loop from OPD identity → charge → pay without WhatsApp screenshots of UPI apps.",
            "Lock patient → add charge → amount due; Show pay QR (mock or test keys); More shows new UI build id "
            "after installing latest share APK",
        ),
        (
            "Era O — Today-only pick, persisted lab orders, browser desk (14–15 Aug 2026)",
            "One picker for all roles; doctor Visit after lock (not after PIN); Lab Desk sees Visit ticks "
            "on another phone; receptionist/staff can work in a browser without an APK.",
            [
                "Remove name/mobile/MRN Select form from Vitals PatientBar; NeedPatient links to /today/#all-patients",
                "pathAfterPatientLock: lab → /labs/, doctor → /visit/ (incl. video), staff → /patient/, receptionist stays /today/",
                "Doctor PIN unlock and app home → /today/ (pick first). Do not bounce doctors off Today if they open it.",
                "Lab nav: Today / Labs / More only; /patient/ redirects lab to Labs if locked else Today; billing never on lab",
                "Directory lock works with MRN only (no phone required); identity endpoint returns clinic_mrn without 10-digit phone",
                "clinic_patients.lab_orders_json + GET/PUT /history/lab-orders (doctor/staff/lab; receptionist 403)",
                "PatientContext loads orders on lock and ~15s history poll; debounce-saves on tick/dismiss",
                "AllPatientsDirectory polls ~15s while visible; WaitingQueue remains a picker for doctor/staff/front desk",
                "Lab tokenize/directory use DoctorSession (not require_clinical) so Lab Desk can lock a patient (RCA-23)",
                "Browser desk: DNS A(+AAAA) app → VPS; nginx serves deploy/web; TLS SAN api+app; CORS includes https://app.*",
                "scripts/deploy_web.cmd builds Next export and publishes to VPS; live https://app.aarogyaoneconnect.in",
                "Share APK v1.48 (versionCode 49) — doctor Today-first; phones still need APK for UI (RCA-22)",
            ],
            "Front desk can run on a PC browser; Lab Desk and Visit stay in sync without websockets; "
            "doctors still pick on Patient Info then land on Visit.",
            "Browser: open https://app.aarogyaoneconnect.in → clinic name + password → receptionist/staff PIN. "
            "Doctor APK: after PIN land on Patient Info; lock from All patients → Visit. Tick a diagnostic on Visit; "
            "Lab Desk on another device sees it under Ordered from Visit within ~15s.",
        ),
    ]

    for title, goal, steps, significance, verify in milestones:
        add_heading(doc, title, 2)
        add_p(doc, f"Goal: {goal}", bold=False)
        add_p(doc, "Development steps:", bold=True, space_after=4)
        add_bullets(doc, steps)
        add_p(doc, f"Significance: {significance}")
        add_p(doc, f"How to verify: {verify}", italic=True)

    add_heading(doc, "Current status snapshot", 2)
    add_table(
        doc,
        ["Area", "Status"],
        [
            ["LAN clinic on PC + phone", "Usable (set clinic URL; not localhost on phone)"],
            ["Roles / PIN / durable sessions", "Implemented"],
            ["MRN / lab / queue / offline / i18n", "Implemented"],
            ["Clinic analytics + CSV export", "Implemented"],
            ["Attachments on disk", "Implemented (legacy base64 still readable)"],
            ["Docker Compose scaffold", "Present; harden before public"],
            ["Public India VPS phones-everywhere", "Documented (CLOUD_DEPLOY.md); not auto-provisioned"],
            ["ABHA / ABDM", "OTP + link implemented; needs NHA credentials for live gateway (ABDM_MOCK for demo)"],
            ["SMS appointments", "Implemented (console/msg91/twilio); day-before remind script"],
            ["Multi-tenant clinics", "Implemented (CLINICS + clinic_id scoping)"],
            ["OPD multi-tab shell", "Implemented (Patient Info / Vitals / Visit / Records / Labs / More)"],
            ["Obstetric profile + GA charts", "Implemented (LMP/EDD/GPLA; Hb; pregnancy X-axis)"],
            ["Case brief / ANC alerts / consult pack", "Implemented (decision support only)"],
            ["Rx pre-sign hints + USG analyze", "Implemented (soft warnings; findings JSON)"],
            ["Demo seed showcase", "scripts/seed_demo.cmd (--wipe)"],
            ["Receptionist role", "Implemented (Patient Info + billing/appointments; no Visit/Rx)"],
            ["Video consult (Jitsi)", "Implemented behind clinic feature video_consult"],
            ["Patient billing ledger", "Implemented (charges/payments + summary on Patient Info)"],
            ["Razorpay UPI QR", "Implemented (env keys + webhook; mock for local; live keys per clinic ops)"],
            ["Hostinger multi-clinic cloud", "Live runbook; dual-stack DNS + nginx restart after API recreate"],
            ["Share HTTPS APK", "scripts/build_share_apk.cmd; More shows UI build stamp (v1.48 as of 14 Aug 2026)"],
            ["Browser clinic desk", "Live at https://app.aarogyaoneconnect.in (same UI; CORS + TLS SAN)"],
            ["Today-only patient pick", "Implemented (All patients / waiting list; no Vitals Select form)"],
            ["Persisted lab orders", "Implemented (lab_orders_json + /history/lab-orders; 15s poll)"],
            ["Play Store listing", "AAB docs + PRIVACY_POLICY.md stub"],
            ["iOS Capacitor project", "Scripts exist; frontend/ios/ not checked in yet"],
            ["Git remote push of Unreleased work", "May still need auth/commit hygiene"],
        ],
    )
    page_break(doc)

    # ========== PART IV ==========
    add_heading(doc, "Part IV — System deep dive", 1)
    add_image(doc, "01_system_architecture.png", "Figure: system architecture")
    add_image(doc, "03_voice_to_rx.png", "Figure: voice to prescription workflow")
    add_image(doc, "04_roles_matrix.png", "Figure: role capabilities")

    add_heading(doc, "9. Data model (conceptual)", 2)
    add_bullets(
        doc,
        [
            "clinical_records — id, clinic_id, blind_patient_id, created_at, encounter_data (JSON).",
            "clinic_sessions — durable unlock tokens (~7 days) including clinic_id.",
            "clinic_queue — today’s waiting list (clinic-scoped).",
            "appointments — scheduled visits; phone encrypted at rest; SMS status fields; modality in_person|video.",
            "payment_intents — Razorpay QR/order state; webhook idempotently writes billing payment rows.",
            "clinic_patients.lab_orders_json — selected/dismissed Visit diagnostics shared with Lab Desk.",
            "encounter_data types: visit/prescription, vitals, document, lab_result, audit, abha_link, "
            "obstetric_profile, billing (charge|payment), video_consult.",
            "Document files live under backend/data/attachments/ (content_path); optional findings JSON "
            "after Analyze; legacy content_base64 still loads.",
            "Existing DBs gain clinic_id / new tables via soft ALTER on API startup (schema_migrate).",
        ],
    )

    add_heading(doc, "10. Identity rules", 2)
    add_bullets(
        doc,
        [
            "Prefer clinic MRN when set → raw key mrn|{MRN} → HMAC.",
            "Else name + 10-digit Indian mobile → name|{digits}.",
            "Phone change remaps history for name|phone keys; MRN-keyed patients skip remap (records_moved=0).",
            "ABHA: local HMAC link always available; with ABDM credentials or ABDM_MOCK, use MOBILE_OTP then link with txn_id.",
            "Linking tokens from ABDM are stored as HMAC only — never cleartext ABHA in columns.",
            "Multi-clinic: every clinical read/write is filtered by the signed-in clinic_id.",
        ],
    )

    add_heading(doc, "11. API map (student overview)", 2)
    add_table(
        doc,
        ["Prefix", "Purpose"],
        [
            ["/health", "Liveness (must stay fast — do not block on Whisper load)"],
            ["/api/v1/auth/*", "Unlock / lock / me / status (users + clinics list)"],
            ["/api/v1/history/*", "Tokenize, search, vitals, docs, lab-results, lab-orders, obstetric-profile, "
             "case-summary, consult-pack, rx-hints, document analyze, billing / billing-summary, change-phone"],
            ["/api/v1/prescription/*", "Transcribe, parse, write PDF, share-link, download"],
            ["/api/v1/queue/*", "Today’s waiting queue"],
            ["/api/v1/appointments/*", "Book / list / remind / cancel / patient-identity (+ SMS)"],
            ["/api/v1/analytics/*", "Today, week, frequency, vitals-trend (+ alerts/lab merge), export.csv"],
            ["/api/v1/integrations/*", "ABHA OTP/link/status, ABDM callback, HL7 ORU + MSA ACK"],
            ["/api/v1/payments/*", "QR create, status poll, Razorpay webhook, mock-pay (dev)"],
            ["/api/v1/video-consult/*", "Mint Jitsi room + join links for locked patient"],
        ],
    )

    add_heading(doc, "12. Frontend map", 2)
    add_bullets(
        doc,
        [
            "app/(clinic)/* — Patient Info (/today) / Vitals (/patient) / Visit / Records / Labs / More "
            "(bottom ClinicNav; receptionist: Patient Info + More; lab: Today / Labs / More)",
            "DoctorGate — clinic picker, PIN, clinic server URL (blocks localhost on phone); doctor PIN → /today/",
            "pathAfterPatientLock (lib/clinicRoutes.ts) — after directory/waiting-list lock: doctor Visit, "
            "staff Vitals, lab Labs, receptionist stays Patient Info",
            "NeedPatient — unlocked Vitals/Labs/Records/Visit prompt + link to /today/#all-patients",
            "PatientBar / PatientContext / LockedPatientChip — lock, MRN, GA chip, ABHA OTP + local link; "
            "no Select form when unlocked; Change patient → All patients",
            "AppointmentScheduler (Google Calendar / .ics + video modality), ClinicAnalytics, WaitingQueue",
            "PatientBilling — charges, payments, amount due, Show pay QR (Razorpay)",
            "ObstetricProfileForm, PatientCaseBrief (alerts, scan cadence, consult pack)",
            "VideoConsultPanel — Jitsi room mint + SMS (feature-gated)",
            "VitalsForm (°C/°F + hemoglobin), PatientVitalsCharts / VitalsTrendCharts, LabResultsForm",
            "EncounterWorkspace (rx-hints before Sign), VoiceRecorder, PrescriptionShare",
            "RecommendedDiagnostics / PendingLabOrders — Visit ticks persist via /history/lab-orders",
            "PatientTimeline, PatientAttachments (Analyze), PatientAuditTrail, OfflineSyncBanner",
            "lib/offlineQueue.ts, lib/i18n.tsx (EN/HI), lib/apiBase.ts, lib/doctorSession.ts, "
            "lib/appBuild.ts (UI build stamp), lib/obstetric.ts, lib/calendarExport.ts",
        ],
    )

    add_heading(doc, "13. Offline semantics", 2)
    add_p(
        doc,
        "If vitals/lab POST fails due to network error or 5xx/408/429, items enqueue in localStorage "
        "(capped). Banner on Today offers Sync now; online event and page load also flush. "
        "Validation failures (HTTP 400/422) are NOT queued — they show the server error and are "
        "dropped from any stale queue on the next flush (see RCA-16). This is not a full offline EHR — "
        "prescriptions still need the API.",
    )

    add_heading(doc, "14. Environment surface (placeholders only)", 2)
    add_p(
        doc,
        "Copy backend/.env.example → backend/.env. Example placeholder lines (DO NOT use these as real secrets):",
    )
    add_bullets(
        doc,
        [
            "SECRET_KEY=CHANGE_ME_GENERATE_WITH_OPENSSL_RAND_HEX_32",
            "SECRET_SALT=CHANGE_ME_GENERATE_WITH_OPENSSL_RAND_HEX_32",
            "CLINIC_USERS=dr1|Dr Example|doctor|pbkdf2$…  (hash via scripts/hash_pin.py)",
            "CLINICS=default|Alpha Clinic|Address|  (optional multi-tenant)",
            "PUBLIC_API_BASE_URL=https://api.your-domain.example",
            "CORS_ORIGINS=…,https://api.your-domain.example,https://app.your-domain.example  (browser desk)",
            "WHISPER_PROVIDER=local|openai|groq ; WHISPER_PRELOAD=false (keep /health fast)",
            "LLM_PROVIDER=ollama|openai|groq",
            "REQUIRE_CLINIC_USERS=true and APP_ENV=production on public hosts",
            "ABDM_MOCK=true (demo OTP 123456) or ABDM_CLIENT_ID / ABDM_CLIENT_SECRET",
            "SMS_PROVIDER=console|msg91|twilio",
            "PAYMENTS_ENABLED=true|false ; RAZORPAY_KEY_ID / KEY_SECRET / WEBHOOK_SECRET ; "
            "RAZORPAY_MOCK=true for local only",
            "VIDEO_CONSULT_PROVIDER=jitsi ; JITSI_BASE_URL=https://meet.jit.si ; "
            "enable video_consult in CLINICS features JSON",
        ],
    )
    add_p(
        doc,
        "Critical student tip: backend/.env.example is a template with empty CLINIC_USERS. "
        "Real users live only in backend/.env (gitignored). Empty example lines do not mean the "
        "clinic has no staff accounts.",
        italic=True,
    )
    page_break(doc)

    # ========== PART V ==========
    add_heading(doc, "Part V — Issues, debugging, and RCA", 1)
    add_image(doc, "06_debug_loop.png", "Figure: student debugging / RCA loop")
    add_p(
        doc,
        "General playbook: (1) write the symptom exactly, (2) list 3–5 hypotheses, (3) gather evidence "
        "(logs, curl, browser Network, adb logcat), (4) change one thing, (5) verify, (6) write the lesson. "
        "Avoid shotgun config edits.",
    )

    rcas = [
        (
            "RCA-1 — uvicorn SettingsError on ALLOWED_HOSTS / CORS_ORIGINS",
            "API refused to start; pydantic-settings error parsing allowed_hosts from .env.",
            "CSV vs JSON; wrong .env path; BOM; validator bug.",
            "Instrument config load; log raw env strings and whether JSON decode ran before validators.",
            "List[str] fields were JSON-decoded before field validators; CSV like localhost,127.0.0.1 failed.",
            "Annotate fields with Annotated[..., NoDecode] so CSV reaches the splitter validator.",
            "Know your settings library’s decode order; .env UX should match how types are parsed.",
        ),
        (
            "RCA-2 — Next.js rejected next.config.ts / App Router broken",
            "Frontend would not start; config.ts unsupported.",
            "Wrong Next major; bad lockfile; React mismatch.",
            "Preflight script logged package.json declared vs installed Next version.",
            "next was pinned to ^9.3.3 (likely npm audit fix --force) while scaffold needed Next 15.",
            "Restore next ^15 and reinstall; avoid audit --force without reading the major bump.",
            "Lockfiles and audit tooling can silently destroy a scaffold.",
        ),
        (
            "RCA-3 — Microphone NotAllowedError on Record",
            "Record failed even on localhost with user gesture.",
            "Permission denied UI; insecure context; AudioContext failure; Permissions-Policy.",
            "Log getUserMedia error name/message and document.permissions.",
            "next.config sent Permissions-Policy: microphone=() which blocks mic.",
            "Change to microphone=(self).",
            "Security headers are feature switches — test mic/camera after every header change.",
        ),
        (
            "RCA-4 — Phone cannot reach API on Wi-Fi",
            "APK or phone browser times out / CORS errors.",
            "uvicorn bound to 127.0.0.1 only; firewall; wrong CORS; wrong PUBLIC_API_BASE_URL.",
            "From phone open http://LAN:8000/health; check Windows Firewall; inspect CORS preflight.",
            "API must listen on 0.0.0.0; CORS must include phone origin; share links need PUBLIC_API_BASE_URL.",
            "Bind 0.0.0.0; update .env CORS and public base; allow ports 8000/3000.",
            "Mobile networking is part of full-stack work.",
        ),
        (
            "RCA-5 — Rebuilding APK for every LAN IP change",
            "Staff moved networks; app pointed at old IP.",
            "Build-time NEXT_PUBLIC_* baked in.",
            "Compare build env vs runtime needs.",
            "Static export baked the API URL at build time.",
            "Runtime “Clinic server address” saved in localStorage (apiBase.ts).",
            "Prefer runtime config for environment-specific endpoints.",
        ),
        (
            "RCA-6 — Sessions vanished after API restart",
            "Doctors unlocked again after every uvicorn reload.",
            "In-memory session dict lost on process exit.",
            "Restart API and observe /auth/me.",
            "Sessions were not durable.",
            "Persist clinic_sessions in SQLite/Postgres with expiry.",
            "Anything you care about across restarts belongs in a database.",
        ),
        (
            "RCA-7 — Open local-doctor mode risk",
            "Empty CLINIC_USERS allowed overly open access in some modes.",
            "Mis-set production flags.",
            "Hit /auth/status in production-like env.",
            "Convenience default unsafe when exposed publicly.",
            "REQUIRE_CLINIC_USERS / APP_ENV=production blocks open mode.",
            "Defaults that help students can harm production — gate them.",
        ),
        (
            "RCA-8 — Whisper missing in Docker image",
            "Voice works on laptop venv but not in stock container.",
            "Image size vs dependency omission.",
            "Read Dockerfile vs requirements.txt.",
            "Dockerfile intentionally skips heavier Whisper deps.",
            "Use WHISPER_PROVIDER=openai|groq on VPS, or enlarge image, or disable voice.",
            "Containers are deliberate subsets — document the tradeoff.",
        ),
        (
            "RCA-9 — Self-signed TLS vs phone trust on 4G",
            "HTTPS with clinic self-signed cert fails off LAN.",
            "Trust store; captive portals; wrong host.",
            "Compare behavior on clinic Wi-Fi vs mobile data.",
            "Self-signed certs are for LAN labs; public phones need Let’s Encrypt.",
            "Public VPS + ACME certificates; point app at https://api.domain.",
            "Match certificate trust model to where phones roam.",
        ),
        (
            "RCA-10 — Collaboration / gitignore surprises",
            "Clone missing android/; cannot push; fear of committing .env.",
            "gitignore; remote auth; large Unreleased delta.",
            "git status; check remotes; confirm .env ignored.",
            "android/ and secrets are intentionally local; remote write access may be limited.",
            "Document generate-android steps; never commit .env; push only with proper access.",
            "Onboarding docs prevent “empty clone” panic.",
        ),
        (
            "RCA-11 — Phone stuck on Loading… / “no profiles” (22 Jul 2026)",
            "APK opened, spinner never finished, or sign-in showed zero users even though CLINIC_USERS seemed set.",
            "API down; wrong clinic URL; empty users; CORS; hung uvicorn still listening on :8000.",
            "From PC: curl http://127.0.0.1:8000/health (must return quickly). From phone: open http://LAN:8000/health. "
            "Compare backend/.env CLINIC_USERS vs .env.example. Check DoctorGate Network call to /auth/status.",
            "Compound causes: (1) API process listening but not answering HTTP (often blocked in Whisper preload/startup); "
            "(2) phone clinic URL left as localhost (points at the phone itself); "
            "(3) confusion between empty .env.example and real .env users (dr1/nurse1/…). "
            "UI previously looked like “no profiles” when the auth fetch failed/hung.",
            "Restart API with WHISPER_PRELOAD=false; set phone Clinic server address to http://192.168.x.x:8000 "
            "(not localhost); keep users in backend/.env; phone first-run now forces clinic URL setup; "
            "auth fetch timeout ~8s with clearer unreachable-server copy; optional scripts/watch_api.cmd.",
            "“Port is open” ≠ “HTTP works.” Always prove /health. Never use localhost as the clinic URL on a physical phone.",
        ),
        (
            "RCA-12 — Android release build fails on ProGuard (AGP 9.3)",
            "Gradle/Capacitor release or AAB build failed looking for proguard-android.txt.",
            "Wrong AGP; missing SDK file; third-party plugin path.",
            "Read Gradle error path; compare Android SDK/AGP version; search plugins for getDefaultProguardFile.",
            "AGP 9.x ships proguard-android-optimize.txt; older getDefaultProguardFile('proguard-android.txt') breaks.",
            "Point app + @capacitor-community/file-opener at optimize file; run frontend/scripts/fix-agp-proguard.mjs "
            "from cap:sync / mobile:build so regenerated android/ stays fixed.",
            "Native toolchain versions drift — pin a fix script when Cap sync regenerates android/.",
        ),
        (
            "RCA-13 — Whisper preload freezes the API at startup",
            "/health hung or took minutes; phones could not unlock; voice later worked only after long wait.",
            "Model download; CPU thrash; uvicorn single worker blocked; firewall.",
            "Time curl /health right after start; watch CPU; read WHISPER_* settings.",
            "WHISPER_PRELOAD=true loaded faster-whisper during lifespan before accepting useful traffic.",
            "Default WHISPER_PRELOAD=false; if true, warm model in a background task after startup; "
            "first transcribe may still be slow — that is preferred over a dead /auth/status.",
            "Never block readiness probes on optional heavy models.",
        ),
        (
            "RCA-14 — Empty CLINIC_USERS in .env.example misread as “no users”",
            "Developer opened .env.example, saw blank CLINIC_USERS, assumed auth was empty.",
            "Wrong file; parser bug; unset env in process.",
            "Open backend/.env (gitignored); GET /api/v1/auth/status; print list_public_users locally.",
            "Template vs secrets file mix-up — real clinic had dr1/nurse1/nurse2/lab1 in .env only.",
            "Document the distinction in handbook/README; UI copy points at .env not .env.example.",
            "Teach two-file mental model: example = placeholders; .env = live secrets.",
        ),
        (
            "RCA-15 — Listen socket up but HTTP never answers",
            "netstat showed :8000 LISTENING; browsers and phones timed out; uvicorn log looked “started.”",
            "Stuck startup hook; deadlock; wrong bind; zombie reloader child.",
            "curl -m 3 /health; kill leftover uvicorn PIDs; restart clean; enable watch_api.cmd.",
            "Process accepted TCP but event loop/startup work prevented request handling (related to RCA-13).",
            "Hard restart; WHISPER_PRELOAD=false; scripts/watch_api.cmd polls /health and restarts after failures.",
            "Operational watchdogs catch “half-alive” APIs that confuse mobile clients.",
        ),
        (
            "RCA-16 — Empty vitals 400 queued as “offline item(s) waiting to sync” (25 Jul 2026)",
            "Today tab showed “1 offline item(s) waiting to sync”; Sync now never cleared it; "
            "uvicorn logged POST /api/v1/history/vitals → 400 Bad Request.",
            "True offline; SECRET_SALT/tokenize fail; range validation; empty payload; UI treating all errors as offline.",
            "Instrument add_vitals_entry + VitalsForm: log which fields were filled, HTTP status, and whether "
            "enqueueOffline ran. Inspect localStorage healthcare_offline_queue. Reproduce empty Save.",
            "Client allowed (or re-synced) an empty vitals body. API correctly returned 400 "
            "(“Enter at least one vital or a diagnostic note”). VitalsForm catch-all treated every !ok "
            "response as offline and enqueued the same empty payload. flushOfflineQueue kept retrying "
            "400s forever, so the Today banner never cleared.",
            "Require at least one vital/note before POST; queue offline only for network/5xx/408/429; "
            "show server detail for 400/422; drop permanent failures on flush; auto-flush on Today mount.",
            "HTTP 400 is not “offline.” Never enqueue validation errors into a sync retry queue.",
        ),
        (
            "RCA-17 — mobile:build fails on Python-style `*` in TypeScript (25 Jul 2026)",
            "npm run mobile:build / next build: Unexpected token `*` in VitalsTrendCharts.tsx seriesFrom.",
            "SWC parser bug; wrong TS target; bad merge; accidental Python syntax in TS.",
            "Read the exact line from the webpack error; compare call sites vs signature.",
            "seriesFrom was declared with Python keyword-only syntax (`*, byGestation: boolean`) while "
            "call sites already passed an options object `{ byGestation }`. Valid Python, illegal TypeScript.",
            "Change signature to `opts: { byGestation: boolean }` and destructure inside. Re-run mobile:build.",
            "Cross-language edits invite syntax bleed — always compile after mixing Python and TypeScript habits.",
        ),
        (
            "RCA-18 — nginx 502 while API /health is OK (Aug 2026)",
            "https://api… returned 502 Bad Gateway; curl to 127.0.0.1:8000/health on the VPS succeeded.",
            "API crash; TLS cert; firewall; wrong compose project; reverse-proxy misconfig.",
            "Compare curl through nginx vs localhost:8000; docker compose ps; docker inspect nginx upstream; "
            "docker compose logs nginx.",
            "After force-recreate api, the API container got a new Docker network IP. nginx had cached the "
            "old upstream address and kept proxying to a dead IP.",
            "docker compose restart nginx (or recreate nginx) after every API force-recreate. Document in "
            "RESTART_PRODUCTION_API.md / CLOUD_DEPLOY.md.",
            "Compose service DNS is not magical forever — restart the proxy when upstream containers move.",
        ),
        (
            "RCA-19 — Some phones / ISPs cannot reach the VPS on IPv4",
            "“Could not reach the clinic server” on mobile data from some networks; SSH/HTTP over IPv4 timed out "
            "from the developer PC while IPv6 worked.",
            "API down; wrong DNS A; firewall; app bug; cert mismatch.",
            "Resolve api hostname (A and AAAA); curl -4 vs curl -6; Hostinger firewall; traceroute.",
            "Public IPv4 path from some India ISPs to the VPS was black-holed; IPv6 path was healthy. "
            "DNS had only an A record, so dual-stack clients still preferred broken IPv4.",
            "Add DNS AAAA for api → VPS IPv6; keep A for IPv4-capable networks; verify /health on both stacks.",
            "Reachability is path-dependent — always test A and AAAA, not only “server is up locally.”",
        ),
        (
            "RCA-20 — Empty nano backend/.env on the VPS (wrong cwd)",
            "Production env “disappeared”; nano opened an empty file; clinic unlock broke after “restoring” env.",
            "Deleted secrets; wrong volume mount; compose env_file path; editor cwd.",
            "pwd; ls -la backend/.env; compare with docker compose config; recover from running container env.",
            "Operator ran nano backend/.env from /root instead of /root/Healthcare-Secure, creating "
            "/root/backend/.env while Compose still mounts /root/Healthcare-Secure/backend/.env.",
            "Always cd /root/Healthcare-Secure first. Prefer recovering live env from the running container "
            "before rewriting. Backup .env before edits.",
            "Path + cwd mistakes look like “secrets vanished.” Confirm absolute paths on every VPS edit.",
        ),
        (
            "RCA-21 — CLINIC_USERS hashes broken after Compose interpolate",
            "Unlock roster missing users or PIN verify failed after putting pbkdf2$… strings into Compose env.",
            "Wrong hash; ROLE typo; REQUIRE_CLINIC_USERS; parser skip.",
            "Log parsed user count (without PINs); compare raw .env vs `docker compose config` output.",
            "Docker Compose treats $ as variable interpolation. Unescaped $ inside pbkdf2$… corrupted hashes.",
            "Write $$ in Compose-managed env values (or keep hashes only in env_file that is not re-interpolated "
            "incorrectly). Verify roster after recreate.",
            "Any $ in secrets used with Compose needs escaping — hashes are not exempt.",
        ),
        (
            "RCA-22 — Phone missing latest UI after “deploying” updates (Aug 2026)",
            "Developer saw Patient Info billing / Show pay QR in source (or laptop browser) but the installed "
            "Android app still looked like the old UI.",
            "Browser cache; service worker; wrong API; role hiding features; Play Store lag; WebView cache.",
            "Compare mtimes: frontend/src vs frontend/out vs android assets vs share/*.apk. Grep built JS for "
            "new strings (e.g. “Show pay QR”). Check More → UI build stamp.",
            "Capacitor embeds a static Next export (webDir=out) inside the APK. Production nginx only proxies "
            "the API — it never serves the React UI. Last share APK (e.g. v1.38, 1 Aug) was built before "
            "billing; source was newer but phones still ran the old bundle.",
            "Rebuild with scripts/build_share_apk.cmd (or mobile:build + assemble); install the new "
            "AarogyaOneConnect-v*.apk; confirm More shows the new UI build id. API recreate alone never "
            "updates screens.",
            "Cloud API deploy ≠ mobile UI deploy. Treat APK rebuild/reinstall as a required release step.",
        ),
        (
            "RCA-23 — Lab Desk 403 when locking a patient from All patients (Aug 2026)",
            "Lab profile could see the directory but tapping a patient failed with HTTP 403; Labs stayed empty.",
            "Wrong PIN/role; feature flag labs off; clinic_id mismatch; tokenize requires phone; staff-only gate.",
            "Reproduce with lab session token; curl tokenize + directory; compare require_clinical vs DoctorSession "
            "on history endpoints; log session.role without PHI.",
            "Tokenize / identity / directory used require_clinical, which treats lab as non-clinical and returns 403. "
            "Lab is a clinical desk for results/uploads but must not bump visit counts like a doctor/staff lock.",
            "Use DoctorSession on tokenize/directory/identity. bump_visit only when session.role != \"lab\". "
            "Keep require_clinical for Rx/sign paths. Tests: lab can GET/PUT lab-orders; receptionist still 403.",
            "Role gates must match the job: “not a doctor” is not the same as “cannot lock a chart.”",
        ),
        (
            "RCA-24 — Browser desk Continue → Failed to fetch (15 Aug 2026)",
            "https://app.aarogyaoneconnect.in clinic name + password Continue showed TypeError Failed to fetch; "
            "PIN screen never appeared.",
            "Wrong baked API URL (localhost); TLS mismatch; API down; CORS missing app origin; empty nginx web root; "
            "stale cached JS; IPv4 black-hole (RCA-19); ad-blocker.",
            "curl OPTIONS/POST /api/v1/auth/clinic-unlock with Origin: https://app…; inspect Access-Control-Allow-Origin; "
            "from the page Runtime fetch status vs localhost:8000; confirm nginx 200 on app host (not 403); "
            "grep built JS for https://api.aarogyaoneconnect.in.",
            "The desk is a static SPA on app.* calling the API on api.* (cross-origin). Until CORS_ORIGINS included "
            "https://app.aarogyaoneconnect.in, the browser hid the CORS failure as Failed to fetch. First nginx "
            "app vhost also 403’d an empty deploy/web before the Next export was published. A hard-refresh after "
            "CORS + static deploy succeeded; fetch(api/auth/status) 200 and clinic-unlock 401 (wrong password) "
            "proved the path — valid credentials then returned the roster.",
            "Add https://app.YOUR_PRODUCT.in to CORS_ORIGINS; recreate API; serve frontend/out via nginx "
            "(scripts/deploy_web.cmd); issue Let’s Encrypt SAN covering api + app; never leave deploy/web empty. "
            "Hard-refresh after publish. Do not confuse this with APK Failed to fetch (that is still RCA-4/19/22).",
            "Browser “Failed to fetch” is often CORS or mixed-content, not a wrong password. Always test the "
            "preflight Origin of the real UI host, not only curl without Origin.",
        ),
    ]

    for title, symptom, hyps, strategy, root, fix, lesson in rcas:
        add_heading(doc, title, 2)
        add_p(doc, f"Symptom: {symptom}")
        add_p(doc, f"Hypotheses considered: {hyps}")
        add_p(doc, f"Debug strategy: {strategy}")
        add_p(doc, f"Root cause: {root}", bold=True)
        add_p(doc, f"Fix: {fix}")
        add_p(doc, f"Lesson: {lesson}", italic=True)

    page_break(doc)

    # ========== PART VI ==========
    add_heading(doc, "Part VI — Secrets, behavior changes, and paid VPS deploy", 1)
    add_p(
        doc,
        "This part is intentionally in scope for the handbook: how to handle secrets safely (without "
        "publishing any), how to change application behavior deliberately, and how to deploy a paid "
        "India-region VPS so phones work everywhere on mobile data.",
    )

    add_heading(doc, "15. Secrets — committing and handling (no confidential values)", 2)
    add_image(doc, "08_secrets_lifecycle.png", "Figure: secrets lifecycle")
    add_p(doc, "What may be committed:", bold=True, space_after=4)
    add_bullets(
        doc,
        [
            ".env.example with CHANGE_ME placeholders",
            "Documentation describing variable names and purpose",
            "hash_pin.py and instructions to generate pbkdf2$ hashes locally",
        ],
    )
    add_p(doc, "What must NEVER be committed:", bold=True, space_after=4)
    add_bullets(
        doc,
        [
            "backend/.env or frontend/.env.local with real values",
            "Real SECRET_KEY / SECRET_SALT / API keys / DB passwords",
            "Plaintext clinic PINs in git history",
            "TLS private keys (deploy/certs/*.key), Play signing keystores (*.jks), key.properties",
        ],
    )
    add_p(doc, "Safe generation examples (run locally; do not paste outputs into chat/docs):", space_after=4)
    add_bullets(
        doc,
        [
            "openssl rand -hex 32   → candidate for SECRET_KEY or SECRET_SALT",
            "python scripts/hash_pin.py   → pbkdf2$… for CLINIC_USERS",
            "Store production values in the VPS .env with chmod restricted to the deploy user",
        ],
    )
    add_p(
        doc,
        "If a secret was accidentally committed: rotate it immediately (new salt remaps patient blinds — "
        "treat as an incident), remove from git history with a deliberate process, and invalidate old sessions.",
    )

    add_heading(doc, "16. Changing application behavior (safe practice)", 2)
    add_p(
        doc,
        "“Changing behavior” means altering features, defaults, validation, or roles. Student-safe process:",
    )
    add_numbered(
        doc,
        [
            "State the user-visible change in one sentence (e.g., “Staff can edit temperature unit”).",
            "Find the owning layer: UI component, API schema, service validation, or env default.",
            "Prefer config flags for environment differences (LAN vs cloud Whisper provider).",
            "Add or extend a pytest when touching auth, identity, or vitals validation.",
            "Update CHANGELOG [Unreleased] and this handbook’s “current status” when behavior ships.",
            "Never “fix” production by editing live DB rows for identity — use official change-phone/MRN flows.",
        ],
    )
    add_p(
        doc,
        "Examples of behavior levers already in the app: WHISPER_PROVIDER, WHISPER_PRELOAD, LLM_PROVIDER, "
        "REQUIRE_CLINIC_USERS, CLINICS / multi-tenant CLINIC_USERS, ABDM_MOCK / ABDM_*, SMS_PROVIDER, "
        "PAYMENTS_ENABLED / RAZORPAY_*, VIDEO_CONSULT_PROVIDER / JITSI_*, CAPACITOR_HTTPS, "
        "vitals age_years and temperature_unit, clinic_mrn vs phone identity, role gates (incl. receptionist), "
        "obstetric LMP/EDD, ANC alert thresholds (code), offline queue permanent-failure drop list.",
    )

    add_heading(doc, "17. Paid VPS deployment (phones everywhere)", 2)
    add_image(doc, "05_cloud_vps_deploy.png", "Figure: India VPS target topology")
    add_p(
        doc,
        "Recommendation: a small always-on VPS in an India region (e.g., Mumbai) costing roughly "
        "Rs 400–1,500/month for API + Postgres + TLS. Free PaaS with cold starts and expiring databases "
        "is unsuitable for live PHI.",
    )

    add_heading(doc, "17.1 Pre-deploy checklist", 3)
    add_numbered(
        doc,
        [
            "Register a domain (example: api.yourclinic.example) — placeholder name only in this doc.",
            "Choose VPS size: 1–2 GB RAM to start; more if you run local Whisper/Ollama.",
            "Decide voice strategy: (A) cloud STT/LLM keys on server, (B) typed notes only, (C) larger box + local models.",
            "Prepare CLINIC_USERS with hashed PINs offline.",
            "Confirm gitignore covers .env, certs, keystores.",
        ],
    )

    add_heading(doc, "17.2 Provision and Docker", 3)
    add_numbered(
        doc,
        [
            "Create VPS in India region; restrict SSH to your IP; open ports 80 and 443 only for the app.",
            "Install Docker Engine + Compose plugin.",
            "Copy release files (without .env) to the server; create .env on the server interactively.",
            "Set DATABASE_URL to Compose Postgres; APP_ENV=production; REQUIRE_CLINIC_USERS=true; "
            "PUBLIC_API_BASE_URL=https://api.yourclinic.example; strong DB password (not the compose default).",
            "Point DNS A record (and AAAA for IPv6) to the VPS public addresses.",
            "Obtain Let’s Encrypt certificates (certbot) covering api and app hosts "
            "(bash deploy/init_letsencrypt.sh api.yourproduct.in app.yourproduct.in) and mount them into nginx.",
            "docker compose up -d --build  (use restart: unless-stopped so reboot brings the stack back)",
            "After any `force-recreate api`, also `docker compose restart nginx` (see RCA-18).",
            "Publish browser desk: scripts\\deploy_web.cmd (Windows) → https://app.yourproduct.in "
            "(CORS must include that origin — see RCA-24).",
            "From a phone on mobile data (not clinic Wi-Fi): open https://api…/health then unlock with PIN.",
            "Ops runbooks: docs/CLOUD_DEPLOY.md and docs/RESTART_PRODUCTION_API.md "
            "(absolute path /root/Healthcare-Secure/backend/.env — see RCA-20).",
        ],
    )

    add_heading(doc, "17.3 Phone APK for cloud", 3)
    add_numbered(
        doc,
        [
            "Preferred share build from repo root: scripts\\build_share_apk.cmd "
            "(sets CAPACITOR_HTTPS=true, bakes NEXT_PUBLIC_API_BASE_URL, copies share\\AarogyaOneConnect-v*.apk).",
            "Or manually: set CAPACITOR_HTTPS=true then npm run mobile:build "
            "(clears cleartext LAN mode — see capacitor.config.ts).",
            "On first run, enter https://api.yourclinic.example as Clinic server address "
            "(phones reject localhost).",
            "After EVERY frontend change, rebuild and reinstall the APK — cloud API restarts do not refresh UI "
            "(RCA-22). Confirm More → UI build stamp. Browser desk is separate: scripts\\deploy_web.cmd "
            "(does not update phones).",
            "Attach privacy policy URL from docs/PRIVACY_POLICY.md before Play listing.",
            "Play Store / signed AAB is optional — see docs/PLAY_STORE.md.",
        ],
    )

    add_heading(doc, "17.4 Backups and operations", 3)
    add_bullets(
        doc,
        [
            "Schedule pg_dump (or scripts/backup_pg.cmd); SQLite backup scripts for laptop DBs.",
            "Store backups encrypted off-box; test restore quarterly.",
            "Monitor GET /health (scripts/watch_api.cmd on LAN; external uptime on VPS — no PHI in probes).",
            "After API force-recreate: restart nginx; verify both localhost:8000 and https://api…/health.",
            "Razorpay: rotate webhook secret with keys; never put RAZORPAY_KEY_SECRET in the APK.",
            "Patch OS/Docker monthly; rotate PINs when staff leave.",
            "Expected monthly cost band: Rs ~400–900 (lean) to Rs ~1,000–2,500 (comfortable + STT usage).",
        ],
    )

    add_heading(doc, "17.5 What this handbook will not do", 3)
    add_p(
        doc,
        "It will not create a cloud account for you, spend money, or print real secrets. "
        "You (or your instructor) provision the VPS; this document is the procedure and rationale.",
    )
    page_break(doc)

    # ========== PART VII ==========
    add_heading(doc, "Part VII — Future enhancement plans", 1)
    add_p(
        doc,
        "Shipped items stay in Part III / current status (through Era O: browser desk, Today-only pick, "
        "persisted lab orders). This table is what remains open for product/ops follow-through.",
    )
    add_table(
        doc,
        ["Theme", "Plan", "Notes"],
        [
            ["Live Razorpay keys on VPS", "Set RAZORPAY_* + webhook; disable RAZORPAY_MOCK", "Code + UI shipped; keys are ops"],
            ["iOS Capacitor project", "cap add ios + TestFlight / Safari static out/", "Android share APK is primary today"],
            ["Live ABDM (not mock)", "NHA sandbox/prod client ID/secret + public callback", "OTP client ready; care-context/HIU still out of scope"],
            ["Richer HL7 / devices", "MLLP framing, more segment types", "ORU + MSA ACK already present"],
            ["Pediatric growth charts UI", "WHO/CDC growth curves by age", "Neonatal bands + adult/pregnancy charts already present"],
            ["MSG91 templates", "DLT-approved SMS templates for India", "console/msg91/twilio + day-before remind wired"],
            ["Play Store publish", "Signed AAB + hosted privacy URL", "docs/PLAY_STORE.md + PRIVACY_POLICY.md"],
            ["Per-clinic letterhead", "Branding from CLINICS rows in PDF", "clinic_id branding helper exists; PDF still mostly global CLINIC_*"],
            ["Richer USG NLP", "OCR + stronger LLM extract when Ollama unavailable", "Heuristic + optional LLM analyze already ships"],
        ],
    )
    page_break(doc)

    # ========== PART VIII ==========
    add_heading(doc, "Part VIII — Appendices", 1)

    add_heading(doc, "A. Glossary", 2)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ["Blind ID", "HMAC hex of normalized patient key"],
            ["MRN", "Clinic medical record number identity key"],
            ["ABHA / ABDM", "Ayushman Bharat Health Account / Digital Mission APIs"],
            ["HIP", "Health Information Provider (clinic facility in ABDM)"],
            ["OPD", "Outpatient department / clinic day"],
            ["ANC", "Antenatal care (pregnancy follow-up)"],
            ["LMP / EDD", "Last menstrual period / estimated due date"],
            ["GPLA", "Gravida / Para / Abortions / Living children"],
            ["Case brief", "Decision-support summary for the locked patient (not a diagnosis)"],
            ["STT", "Speech-to-text"],
            ["AAB", "Android App Bundle for Play Store"],
            ["DPDP", "India Digital Personal Data Protection Act (compliance context)"],
            ["RCA", "Root cause analysis"],
            ["clinic_id", "Multi-tenant clinic key on sessions and records"],
            ["Patient Info", "Nav label for /today — OPD ops + billing (was “Today”)"],
            ["Receptionist", "Front-desk role: Patient Info + More; no Visit/Rx sign"],
            ["Payment intent", "Razorpay QR/order row awaiting webhook → ledger payment"],
            ["UPI QR", "Scan-to-pay QR shown from Patient Billing"],
            ["UI build stamp", "APP_BUILD_ID shown on More — identifies Capacitor static bundle"],
            ["AAAA", "DNS record for IPv6 address of the API host"],
            ["Browser desk", "https://app.aarogyaoneconnect.in — static clinic UI for PC/browser (esp. front desk/staff)"],
            ["lab orders", "Visit recommended-diagnostics ticks persisted on the patient for Lab Desk"],
            ["pathAfterPatientLock", "Role-based route after locking from All patients or waiting list"],
        ],
    )

    add_heading(doc, "B. Command cheat sheet (Windows)", 2)
    add_bullets(
        doc,
        [
            "Backend venv: cd backend && python -m venv .venv && .\\.venv\\Scripts\\pip install -r requirements.txt",
            "API: .\\.venv\\Scripts\\uvicorn app.main:app --reload --host 0.0.0.0 --port 8000",
            "Prove API alive: curl -m 3 http://127.0.0.1:8000/health",
            "Watchdog: scripts\\watch_api.cmd",
            "Demo seed: scripts\\seed_demo.cmd [--wipe]",
            "Day-before SMS: scripts\\remind_day_before.cmd",
            "Frontend: cd frontend && npm install && npm run dev",
            "Mobile (LAN): npm run mobile:build && npm run cap:open:android",
            "Mobile (HTTPS share APK): scripts\\build_share_apk.cmd → share\\AarogyaOneConnect-v*.apk",
            "Browser desk publish: scripts\\deploy_web.cmd → https://app.aarogyaoneconnect.in",
            "Mobile (HTTPS manual): set CAPACITOR_HTTPS=true && npm run mobile:build",
            "Tests: cd backend && .\\.venv\\Scripts\\python -m pytest -q",
            "PIN hash: .\\.venv\\Scripts\\python ..\\scripts\\hash_pin.py",
            "Docker: docker compose up --build",
            "After API recreate on VPS: docker compose restart nginx",
            "Production API restart runbook: docs\\RESTART_PRODUCTION_API.md",
            "Regenerate handbook diagrams: python docs/build_handbook_diagrams.py",
            "Rebuild this Word doc: python docs/build_handbook_docx.py",
        ],
    )

    add_heading(doc, "C. Annotated top-level tree", 2)
    add_bullets(
        doc,
        [
            "backend/ — FastAPI API, models, services, tests, Dockerfile",
            "frontend/ — Next.js UI, Capacitor config, components",
            "deploy/ — nginx TLS config, cert generator (certs gitignored)",
            "docs/ — handbook (.docx), CLOUD_DEPLOY, RESTART_PRODUCTION_API, PLAY_STORE, "
            "PRIVACY_POLICY, ABDM_SMS_MULTI_TENANT, DEMO_CLIENT",
            "scripts/ — backup/restore, watch_api, seed_demo, remind_day_before, build_share_apk, "
            "deploy_web, Ollama helpers, AAB, hash_pin",
            "share/ — generated HTTPS APKs for clinic phones (do not commit secrets beside them)",
            "deploy/web — published Next static export for the browser desk (gitignored except .gitkeep)",
            "docker-compose.yml — Postgres + API + nginx",
            "CHANGELOG.md — SemVer history",
        ],
    )

    add_heading(doc, "D. How to keep this handbook as the single source of truth", 2)
    add_numbered(
        doc,
        [
            "When you ship a user-visible change, update CHANGELOG [Unreleased] and rebuild this .docx.",
            "Add new RCA entries when a non-obvious bug is fixed.",
            "Never paste production secrets into the handbook source or screenshots.",
            "Prefer editing docs/build_handbook_docx.py (or regenerating from an approved outline) over hand-editing the .docx forever.",
        ],
    )

    add_heading(doc, "E. Closing note for students", 2)
    add_p(
        doc,
        "You now have one document that explains why the app exists, how it was built, how Android packaging "
        "and the browser desk fit, how patient identity, billing, light analytics, and gynae decision support work, "
        "what broke and why (including cloud nginx/DNS, stale-APK, Lab Desk 403, and CORS Failed-to-fetch lessons), "
        "how to handle secrets, how to change behavior carefully, and how to host on a paid India VPS for phones "
        "and PCs used everywhere. Build small, verify often, rebuild the APK when the phone UI changes, "
        "redeploy the web desk when the browser UI changes, and treat clinic data with respect.",
    )

    try:
        doc.save(OUT)
        print("Wrote", OUT)
        return OUT
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print("Primary .docx locked; wrote", OUT_FALLBACK)
        print("Close Word and re-run to overwrite Healthcare_Secure_Handbook.docx")
        return OUT_FALLBACK


if __name__ == "__main__":
    build()
